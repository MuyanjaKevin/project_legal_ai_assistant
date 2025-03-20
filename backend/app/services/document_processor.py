import os
import PyPDF2
from docx import Document
import io
from app.config.database import get_database
from typing import Optional

class DocumentProcessor:
    def __init__(self):
        self.supported_types = ['pdf', 'docx', 'txt']

    def extract_text(self, content: bytes) -> str:
        """Extract text from document content"""
        if isinstance(content, str):
            return content
            
        if isinstance(content, bytes):
            try:
                # Try PDF
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
            except:
                try:
                    # Try DOCX
                    doc = Document(io.BytesIO(content))
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return text
                except:
                    # Return raw text if other formats fail
                    return content.decode('utf-8', errors='ignore')
        
        return ""

    def preprocess_text(self, text: str) -> str:
        """Clean and prepare text for comparison"""
        if not text:
            return ""
        # Remove extra whitespace
        text = " ".join(text.split())
        # Convert to lowercase for better comparison
        return text.lower()

def extract_text_from_document(document_id) -> Optional[str]:
    """Extract text from document based on its file type"""
    try:
        db = get_database()
        
        # Get document from database
        document = db.documents.find_one({"_id": document_id})
        if not document:
            return None
        
        file_path = document["file_path"]
        file_type = document["file_type"].lower()
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extracted_text = ""
        
        # Extract text based on file type
        if file_type == "pdf":
            extracted_text = extract_from_pdf(file_path)
        elif file_type == "docx":
            extracted_text = extract_from_docx(file_path)
        elif file_type == "txt":
            extracted_text = extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Update document with extracted text
        if extracted_text:
            db.documents.update_one(
                {"_id": document_id},
                {"$set": {"extracted_text": extracted_text, "text_extracted": True}}
            )
        
        return extracted_text
    except Exception as e:
        print(f"Error processing document {document_id}: {str(e)}")
        return None

def extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        raise
    return text

def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        raise
    return text

def extract_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try alternative encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            raise
    except Exception as e:
        print(f"Error extracting text from TXT: {e}")
        raise